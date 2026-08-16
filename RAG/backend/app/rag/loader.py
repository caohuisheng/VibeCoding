"""文档解析：PDF / Word(.docx) / Markdown / TXT → LangChain Document。"""

import io

from langchain_core.documents import Document


def load_document(file_bytes: bytes, filename: str) -> list[Document]:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        return _load_pdf(file_bytes, filename)
    if ext == "docx":
        return _load_docx(file_bytes, filename)
    if ext in ("md", "markdown", "txt", "text"):
        return _load_text(file_bytes, filename)
    raise ValueError(f"不支持的文件类型: .{ext}")


def _load_pdf(data: bytes, filename: str) -> list[Document]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(
                Document(page_content=text, metadata={"source": filename, "page": i + 1})
            )
    return docs


def _load_docx(data: bytes, filename: str) -> list[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    if not text.strip():
        return []
    return [Document(page_content=text, metadata={"source": filename, "page": 1})]


def _load_text(data: bytes, filename: str) -> list[Document]:
    text = data.decode("utf-8", errors="ignore")
    if not text.strip():
        return []
    return [Document(page_content=text, metadata={"source": filename, "page": 1})]
